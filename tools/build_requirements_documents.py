from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
FRAME_DIR = Path("/private/tmp/personal-task-concept-frames")
SKILL_DIR = Path(
    "/Users/haruki.shimo/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.723.12215/skills/documents"
)

GREEN = "0D6B57"
DARK_GREEN = "19322F"
MUTED = "66736E"
LIGHT_GREEN = "E4F0EB"
LIGHT_GRAY = "F2F4F3"
BORDER = "D7DEDA"
WHITE = "FFFFFF"
BLACK = "202724"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def font_path() -> str:
    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("No suitable CJK font found")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    # The installed CJK collection renders Japanese reliably for diagram labels.
    return ImageFont.truetype(font_path(), size=size, index=0)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str | None = None,
    radius: int = 26,
    width: int = 2,
) -> None:
    draw.rounded_rectangle(
        xy,
        radius=radius,
        fill=fill,
        outline=outline,
        width=width,
    )


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str,
    width: int = 5,
) -> None:
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    draw.polygon(
        [(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)],
        fill=color,
    )


def create_screen_map() -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    destination = ASSETS_DIR / "screen-map.png"
    canvas = Image.new("RGB", (1400, 680), "#F1F0EA")
    draw = ImageDraw.Draw(canvas)
    draw.text(
        (55, 30),
        "主要4画面",
        fill=f"#{DARK_GREEN}",
        font=pil_font(34, bold=True),
    )
    draw.text(
        (260, 42),
        "期限と優先度だけに集中する情報設計",
        fill=f"#{MUTED}",
        font=pil_font(18),
    )

    screens = [
        ("01  ダッシュボード", FRAME_DIR / "frame-0105.png"),
        ("02  TODO ALL", FRAME_DIR / "frame-0253.png"),
        ("03  今日まで", FRAME_DIR / "frame-0325.png"),
        ("04  優先度マトリクス", FRAME_DIR / "frame-0411.png"),
    ]
    positions = [70, 395, 720, 1045]

    for (label, source), x in zip(screens, positions):
        rounded_box(
            draw,
            (x - 12, 91, x + 282, 632),
            "#FBFCFA",
            "#D9DFDB",
            radius=24,
            width=2,
        )
        if source.exists():
            frame = Image.open(source).convert("RGB")
            phone = frame.crop((426, 20, 854, 880))
            phone.thumbnail((270, 500), Image.Resampling.LANCZOS)
            canvas.paste(phone, (x, 106))
        draw.text(
            (x + 135, 642),
            label,
            fill=f"#{GREEN}",
            font=pil_font(17, bold=True),
            anchor="mm",
        )

    canvas.save(destination, quality=95)
    return destination


def create_architecture_diagram() -> Path:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    destination = ASSETS_DIR / "architecture.png"
    canvas = Image.new("RGB", (1400, 610), "#F7F8F6")
    draw = ImageDraw.Draw(canvas)

    draw.text(
        (55, 35),
        "システム構成",
        fill=f"#{DARK_GREEN}",
        font=pil_font(36, bold=True),
    )
    draw.text(
        (55, 85),
        "ブラウザに秘密情報を渡さず、Vercel上のNext.jsを唯一のAPI窓口にする",
        fill=f"#{MUTED}",
        font=pil_font(20),
    )

    rounded_box(draw, (55, 195, 305, 455), "#FFFFFF", "#D7DEDA", 28, 3)
    draw.text(
        (180, 245),
        "ブラウザ",
        fill=f"#{DARK_GREEN}",
        font=pil_font(28, bold=True),
        anchor="mm",
    )
    draw.text(
        (180, 305),
        "PC / スマートフォン",
        fill=f"#{MUTED}",
        font=pil_font(18),
        anchor="mm",
    )
    draw.text(
        (180, 355),
        "署名済みCookie",
        fill=f"#{GREEN}",
        font=pil_font(18, bold=True),
        anchor="mm",
    )
    draw.text(
        (180, 390),
        "タスク画面のみ",
        fill=f"#{MUTED}",
        font=pil_font(16),
        anchor="mm",
    )

    rounded_box(draw, (410, 155, 815, 495), "#E4F0EB", "#0D6B57", 30, 4)
    draw.text(
        (612, 205),
        "Next.js on Vercel",
        fill=f"#{GREEN}",
        font=pil_font(30, bold=True),
        anchor="mm",
    )
    rounded_box(draw, (450, 252, 775, 324), "#FFFFFF", "#C9D8D2", 18, 2)
    draw.text(
        (612, 288),
        "App Router / 4画面",
        fill=f"#{DARK_GREEN}",
        font=pil_font(20, bold=True),
        anchor="mm",
    )
    rounded_box(draw, (450, 344, 775, 416), "#FFFFFF", "#C9D8D2", 18, 2)
    draw.text(
        (612, 380),
        "認証・Route Handlers",
        fill=f"#{DARK_GREEN}",
        font=pil_font(20, bold=True),
        anchor="mm",
    )
    draw.text(
        (612, 458),
        "入力検証 / 秘密情報 / API変換",
        fill=f"#{MUTED}",
        font=pil_font(16),
        anchor="mm",
    )

    rounded_box(draw, (925, 195, 1345, 455), "#FFFFFF", "#D7DEDA", 28, 3)
    draw.text(
        (1135, 245),
        "Google Sheets API v4",
        fill=f"#{DARK_GREEN}",
        font=pil_font(25, bold=True),
        anchor="mm",
    )
    rounded_box(draw, (985, 298, 1285, 385), "#FCE9E5", "#E6B9B1", 18, 2)
    draw.text(
        (1135, 330),
        "非公開 Tasks シート",
        fill="#A74236",
        font=pil_font(19, bold=True),
        anchor="mm",
    )
    draw.text(
        (1135, 363),
        "サービスアカウントだけに共有",
        fill=f"#{MUTED}",
        font=pil_font(14),
        anchor="mm",
    )

    draw_arrow(draw, (305, 325), (410, 325), f"#{GREEN}", 5)
    draw.text(
        (357, 294),
        "HTTPS",
        fill=f"#{MUTED}",
        font=pil_font(14, bold=True),
        anchor="mm",
    )
    draw_arrow(draw, (815, 325), (925, 325), f"#{GREEN}", 5)
    draw.text(
        (870, 294),
        "認証済みAPI",
        fill=f"#{MUTED}",
        font=pil_font(14, bold=True),
        anchor="mm",
    )

    draw.text(
        (700, 560),
        "APP_PASSPHRASE_HASH・SESSION_SECRET・Google秘密鍵はVercel環境変数で管理",
        fill=f"#{MUTED}",
        font=pil_font(17),
        anchor="mm",
    )
    canvas.save(destination, quality=95)
    return destination


def set_run_font(
    run,
    *,
    name: str = "Noto Sans CJK JP",
    east_asia: str = "Noto Sans CJK JP",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_value in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_value))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def choose_widths(rows: list[list[str]]) -> list[int]:
    column_count = len(rows[0])
    max_lengths = []
    for index in range(column_count):
        max_lengths.append(
            max(
                3,
                max(
                    len(re.sub(r"[`*_]", "", row[index]))
                    for row in rows
                    if index < len(row)
                ),
            )
        )
    if column_count == 2:
        first = min(3000, max(1600, int(PAGE_WIDTH_DXA * max_lengths[0] / sum(max_lengths))))
        return [first, PAGE_WIDTH_DXA - first]
    if column_count == 3:
        if max_lengths[0] <= 6:
            return [1100, 2550, 5710]
        return [2200, 3100, 4060]
    if column_count == 4:
        if rows[0][0] == "変数":
            return [2700, 1600, 1900, 3160]
        if rows[0][0] == "ID":
            return [1100, 2600, 3800, 1860]
        return [1500, 1200, 1800, 4860]
    if column_count == 5:
        return [900, 1100, 1100, 1200, 5060]
    equal = PAGE_WIDTH_DXA // column_count
    widths = [equal] * column_count
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_hyperlink(paragraph, text_value: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text_value
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|https?://\S+)"
)


def add_inline(paragraph, source: str, *, size: float = 11, color: str = BLACK) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(source):
        if match.start() > position:
            run = paragraph.add_run(source[position : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                name="SF Mono",
                east_asia="Noto Sans CJK JP",
                size=max(8.5, size - 1),
                color=DARK_GREEN,
            )
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EDF2EF")
            run._element.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            add_hyperlink(paragraph, token, token.rstrip(".,"))
        position = match.end()
    if position < len(source):
        run = paragraph.add_run(source[position:])
        set_run_font(run, size=size, color=color)


def shade_paragraph(paragraph, fill: str, border_color: str | None = None) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    if border_color:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "16")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)
        ppr.append(p_bdr)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_numbering_instance(
    doc: Document,
    start: int = 1,
    *,
    compact_reference: bool = False,
) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), "1")
    level.append(start_node)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "269" if compact_reference else "360")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540" if compact_reference else "720")
    indent.set(qn("w:hanging"), "271" if compact_reference else "360")
    ppr.append(indent)
    level.append(ppr)
    abstract_num.append(level)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(
    paragraph,
    num_id: int,
    *,
    compact_reference: bool = False,
) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    paragraph.paragraph_format.left_indent = Inches(
        0.375 if compact_reference else 0.5
    )
    paragraph.paragraph_format.first_line_indent = Inches(
        -0.188 if compact_reference else -0.25
    )
    paragraph.paragraph_format.space_after = Pt(4 if compact_reference else 8)
    paragraph.paragraph_format.line_spacing = (
        1.25 if compact_reference else 1.167
    )


def add_table(
    doc: Document,
    rows: list[list[str]],
    *,
    compact_reference: bool = False,
) -> None:
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, choose_widths(rows))
    set_table_borders(table)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(
                paragraph,
                value,
                size=9.3 if len(rows[0]) >= 4 else 9.8,
                color=DARK_GREEN if row_index == 0 else BLACK,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                set_cell_shading(
                    cells[index],
                    "E8EEF5" if compact_reference else LIGHT_GREEN,
                )
            elif row_index % 2 == 0:
                set_cell_shading(cells[index], "F8F9F8")
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
        if row_index == 0:
            set_repeat_table_header(table.rows[-1])
    set_table_geometry(table, choose_widths(rows))

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(
    doc: Document,
    running_title: str,
    *,
    compact_reference: bool = False,
) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1 if compact_reference else 0.85)
    section.bottom_margin = Inches(1 if compact_reference else 0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492 if compact_reference else 0.42)
    section.footer_distance = Inches(0.492 if compact_reference else 0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK JP"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK JP")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK JP")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if compact_reference else 1.1

    if compact_reference:
        style_map = {
            "Heading 1": (16, GREEN, 18, 10),
            "Heading 2": (13, GREEN, 14, 7),
            "Heading 3": (12, DARK_GREEN, 10, 5),
        }
    else:
        style_map = {
            "Heading 1": (16, GREEN, 16, 8),
            "Heading 2": (13, GREEN, 12, 6),
            "Heading 3": (11.5, DARK_GREEN, 8, 4),
        }
    for name, (size, color, before, after) in style_map.items():
        style = doc.styles[name]
        style.font.name = "Noto Sans CJK JP"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Noto Sans CJK JP"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
        style.paragraph_format.left_indent = Inches(
            0.375 if compact_reference else 0.5
        )
        style.paragraph_format.first_line_indent = Inches(
            -0.188 if compact_reference else -0.25
        )
        style.paragraph_format.space_after = Pt(4 if compact_reference else 8)
        style.paragraph_format.line_spacing = (
            1.25 if compact_reference else 1.167
        )

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(running_title)
    set_run_font(run, size=9, color=MUTED, bold=True)
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.35))
    run = hp.add_run("\tDraft · 2026-07-27")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("わたしのタスク管理  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(fp)

    doc.core_properties.author = "Codex"
    doc.core_properties.subject = "個人専用タスク管理アプリ"
    doc.core_properties.keywords = "Next.js, Vercel, Google Sheets, タスク管理"


def add_title_block(doc: Document, title_value: str, subtitle_value: str) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title_value)
    set_run_font(run, size=25, color=DARK_GREEN, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle_value)
    set_run_font(run, size=14, color=GREEN, bold=True)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in raw):
            rows.append(raw)
        index += 1
    return rows, index


def build_docx(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    compact_reference = (
        "実装手順書" in markdown_path.name
        or "AIエージェント一覧" in markdown_path.name
    )
    if "AIエージェント一覧" in markdown_path.name:
        running_title = "AIエージェント一覧"
    elif compact_reference:
        running_title = "実装手順書"
    elif "技術" in markdown_path.name:
        running_title = "技術要件書"
    else:
        running_title = "要件定義書"
    configure_document(
        doc,
        running_title,
        compact_reference=compact_reference,
    )

    first_title = next(line[2:] for line in lines if line.startswith("# "))
    first_subtitle = next(line[3:] for line in lines if line.startswith("## "))
    add_title_block(doc, first_title, first_subtitle)

    index = 0
    title_seen = False
    subtitle_seen = False
    in_code = False
    code_lines: list[str] = []
    front_matter = True
    active_num_id: int | None = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        number_match = re.match(r"(\d+)\. (.+)", stripped)
        if not in_code and number_match is None:
            active_num_id = None

        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.right_indent = Inches(0.18)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(9)
                paragraph.paragraph_format.line_spacing = 1.0
                shade_paragraph(paragraph, "F3F5F4", BORDER)
                run = paragraph.add_run("\n".join(code_lines))
                set_run_font(
                    run,
                    name="SF Mono",
                    east_asia="Noto Sans CJK JP",
                    size=8.6,
                    color=DARK_GREEN,
                )
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped == "---":
            front_matter = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(14)
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "10")
            bottom.set(qn("w:space"), "4")
            bottom.set(qn("w:color"), GREEN)
            pbdr.append(bottom)
            ppr.append(pbdr)
            index += 1
            continue

        if raw.startswith("# "):
            if not title_seen:
                title_seen = True
            else:
                paragraph = doc.add_paragraph(style="Heading 1")
                add_inline(paragraph, raw[2:], size=16, color=GREEN)
            index += 1
            continue

        if raw.startswith("## "):
            if not subtitle_seen:
                subtitle_seen = True
            else:
                paragraph = doc.add_paragraph(style="Heading 1")
                add_inline(paragraph, raw[3:], size=16, color=GREEN)
            index += 1
            continue

        if raw.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, raw[4:], size=13, color=GREEN)
            index += 1
            continue

        if raw.startswith("#### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, raw[5:], size=11.5, color=DARK_GREEN)
            index += 1
            continue

        if stripped.startswith("!["):
            match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if match:
                alt, relative_path = match.groups()
                image_path = (markdown_path.parent / relative_path).resolve()
                if image_path.exists():
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(4)
                    run = paragraph.add_run()
                    shape = run.add_picture(str(image_path), width=Inches(6.3))
                    doc_pr = shape._inline.docPr
                    doc_pr.set("descr", alt)
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(10)
                    run = caption.add_run(alt)
                    set_run_font(run, size=9, color=MUTED, italic=True)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            if rows:
                add_table(
                    doc,
                    rows,
                    compact_reference=compact_reference,
                )
            continue

        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(
                10 if compact_reference else 12
            )
            paragraph.paragraph_format.line_spacing = (
                1.25 if compact_reference else 1.15
            )
            shade_paragraph(paragraph, LIGHT_GREEN, GREEN)
            add_inline(paragraph, stripped.lstrip("> "), size=11, color=DARK_GREEN)
            index += 1
            continue

        checklist_match = re.match(r"- \[([ xX])\] (.+)", stripped)
        if checklist_match:
            checked, value = checklist_match.groups()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.22)
            paragraph.paragraph_format.space_after = Pt(
                4 if compact_reference else 5
            )
            paragraph.paragraph_format.line_spacing = (
                1.25 if compact_reference else 1.1
            )
            run = paragraph.add_run("☒ " if checked.lower() == "x" else "☐ ")
            set_run_font(run, size=11, color=GREEN, bold=True)
            add_inline(paragraph, value, size=10.5, color=BLACK)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:], size=11, color=BLACK)
            index += 1
            continue

        if number_match:
            if active_num_id is None:
                active_num_id = create_numbering_instance(
                    doc,
                    start=int(number_match.group(1)),
                    compact_reference=compact_reference,
                )
            paragraph = doc.add_paragraph()
            apply_numbering(
                paragraph,
                active_num_id,
                compact_reference=compact_reference,
            )
            add_inline(paragraph, number_match.group(2), size=11, color=BLACK)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6 if not front_matter else 3)
        paragraph.paragraph_format.line_spacing = (
            1.25 if compact_reference else 1.1
        )
        add_inline(
            paragraph,
            stripped,
            size=10.5 if front_matter else 11,
            color=MUTED if front_matter else BLACK,
        )
        index += 1

    doc.core_properties.title = first_title
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    create_screen_map()
    create_architecture_diagram()
    pairs = [
        (
            DOCS_DIR / "要件定義書_わたしのタスク管理_v0.1.md",
            DOCS_DIR / "要件定義書_わたしのタスク管理_v0.1.docx",
        ),
        (
            DOCS_DIR / "技術要件書_わたしのタスク管理_v0.1.md",
            DOCS_DIR / "技術要件書_わたしのタスク管理_v0.1.docx",
        ),
        (
            DOCS_DIR / "実装手順書_わたしのタスク管理_v0.1.md",
            DOCS_DIR / "実装手順書_わたしのタスク管理_v0.1.docx",
        ),
        (
            DOCS_DIR / "AIエージェント一覧_わたしのタスク管理_v0.1.md",
            DOCS_DIR / "AIエージェント一覧_わたしのタスク管理_v0.1.docx",
        ),
    ]
    for markdown_path, output_path in pairs:
        build_docx(markdown_path, output_path)
        print(output_path)


if __name__ == "__main__":
    main()
